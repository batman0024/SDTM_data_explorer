
import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document

def export_xpt(df: pd.DataFrame, dataset_name: str = 'DATASET', suggested_name: str = 'dataset.xpt'):
    csv = df.to_csv(index=False).encode('utf-8')
    st.download_button("⬇️ Download CSV (XPT alt)", data=csv, file_name=suggested_name.replace('.xpt', '.csv'), mime='text/csv')

def export_table_to_excel(df: pd.DataFrame, sheet_name: str = 'TFL', suggested_name: str = 'table.xlsx'):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name=sheet_name)
    st.download_button("⬇️ Download Excel", data=output.getvalue(), file_name=suggested_name, mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

def export_table_to_word(df: pd.DataFrame, title: str = 'Table', suggested_name: str = 'table.docx'):
    output = BytesIO()
    doc = Document(); doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns): hdr_cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row): cells[i].text = '' if pd.isna(val) else str(val)
    doc.save(output)
    st.download_button("⬇️ Download Word", data=output.getvalue(), file_name=suggested_name, mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
